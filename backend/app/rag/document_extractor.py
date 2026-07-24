from __future__ import annotations

import csv
import json
import logging
import re
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import HTTPException, status
from markdown import markdown
from pypdf import PdfReader

logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".txt",
    ".md",
    ".docx",
    ".csv",
    ".json",
}


class DocumentExtractionError(Exception):
    """
    Raised when text cannot be extracted from a document.
    """


def clean_extracted_text(text: str) -> str:
    """
    Clean extracted document text.

    This function:
    - removes null characters
    - normalizes Windows and Mac line endings
    - removes excessive spaces
    - removes excessive blank lines
    """

    if not text:
        return ""

    cleaned_text = text.replace("\x00", "")

    cleaned_text = cleaned_text.replace(
        "\r\n",
        "\n",
    )

    cleaned_text = cleaned_text.replace(
        "\r",
        "\n",
    )

    cleaned_text = re.sub(
        r"[ \t]+",
        " ",
        cleaned_text,
    )

    cleaned_text = re.sub(
        r" *\n *",
        "\n",
        cleaned_text,
    )

    cleaned_text = re.sub(
        r"\n{3,}",
        "\n\n",
        cleaned_text,
    )

    return cleaned_text.strip()


def extract_text_from_pdf(
    file_path: Path,
) -> str:
    """
    Extract text from a PDF file using pypdf.
    """

    try:
        reader = PdfReader(str(file_path))

        extracted_pages: list[str] = []

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            try:
                page_text = page.extract_text() or ""

                if page_text.strip():
                    extracted_pages.append(f"[Page {page_number}]\n{page_text}")

            except Exception:
                logger.exception(
                    "Unable to extract PDF page %s from %s",
                    page_number,
                    file_path,
                )

        return clean_extracted_text("\n\n".join(extracted_pages))

    except Exception as exc:
        logger.exception(
            "PDF extraction failed for %s",
            file_path,
        )

        raise DocumentExtractionError(f"Unable to extract text from PDF: {exc}") from exc


def extract_pdf_pages(
    file_path: Path,
) -> list[dict[str, Any]]:
    """
    Extract PDF text page by page.

    Returns:
        [
            {
                "page_number": 1,
                "text": "Page content"
            }
        ]
    """

    try:
        reader = PdfReader(str(file_path))

        pages: list[dict[str, Any]] = []

        for page_number, page in enumerate(
            reader.pages,
            start=1,
        ):
            try:
                page_text = clean_extracted_text(page.extract_text() or "")

                if page_text:
                    pages.append(
                        {
                            "page_number": page_number,
                            "text": page_text,
                        }
                    )

            except Exception:
                logger.exception(
                    "Unable to extract page %s from PDF %s",
                    page_number,
                    file_path,
                )

        return pages

    except Exception as exc:
        logger.exception(
            "PDF page extraction failed for %s",
            file_path,
        )

        raise DocumentExtractionError(f"Unable to extract PDF pages: {exc}") from exc


def extract_text_from_txt(
    file_path: Path,
) -> str:
    """
    Extract text from TXT files.

    Attempts multiple common encodings.
    """

    encodings = [
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
    ]

    last_error: Exception | None = None

    for encoding in encodings:
        try:
            text = file_path.read_text(encoding=encoding)

            return clean_extracted_text(text)

        except UnicodeDecodeError as exc:
            last_error = exc

    raise DocumentExtractionError(f"Unable to decode text file: {last_error}")


def extract_text_from_markdown(
    file_path: Path,
) -> str:
    """
    Extract plain text from Markdown.
    """

    try:
        markdown_text = extract_text_from_txt(file_path)

        html_content = markdown(markdown_text)

        soup = BeautifulSoup(
            html_content,
            "html.parser",
        )

        plain_text = soup.get_text(separator="\n")

        return clean_extracted_text(plain_text)

    except Exception as exc:
        logger.exception(
            "Markdown extraction failed for %s",
            file_path,
        )

        raise DocumentExtractionError(f"Unable to extract Markdown text: {exc}") from exc


def extract_text_from_docx(
    file_path: Path,
) -> str:
    """
    Extract paragraphs and tables from a DOCX file.
    """

    try:
        document = DocxDocument(str(file_path))

        content: list[str] = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                content.append(paragraph_text)

        for table_number, table in enumerate(
            document.tables,
            start=1,
        ):
            table_rows: list[str] = []

            for row in table.rows:
                row_values = [clean_extracted_text(cell.text) for cell in row.cells]

                row_values = [value for value in row_values if value]

                if row_values:
                    table_rows.append(" | ".join(row_values))

            if table_rows:
                content.append(f"[Table {table_number}]")

                content.extend(table_rows)

        return clean_extracted_text("\n\n".join(content))

    except Exception as exc:
        logger.exception(
            "DOCX extraction failed for %s",
            file_path,
        )

        raise DocumentExtractionError(f"Unable to extract DOCX text: {exc}") from exc


def extract_text_from_csv(
    file_path: Path,
) -> str:
    """
    Extract CSV rows as readable plain text.
    """

    encodings = [
        "utf-8",
        "utf-8-sig",
        "cp1252",
        "latin-1",
    ]

    last_error: Exception | None = None

    for encoding in encodings:
        try:
            rows: list[str] = []

            with file_path.open(
                mode="r",
                encoding=encoding,
                newline="",
            ) as csv_file:
                reader = csv.reader(csv_file)

                for row in reader:
                    cleaned_row = [clean_extracted_text(value) for value in row]

                    rows.append(" | ".join(cleaned_row))

            return clean_extracted_text("\n".join(rows))

        except UnicodeDecodeError as exc:
            last_error = exc

        except Exception as exc:
            logger.exception(
                "CSV extraction failed for %s",
                file_path,
            )

            raise DocumentExtractionError(f"Unable to extract CSV text: {exc}") from exc

    raise DocumentExtractionError(f"Unable to decode CSV file: {last_error}")


def extract_text_from_json(
    file_path: Path,
) -> str:
    """
    Extract JSON as formatted readable text.
    """

    try:
        raw_text = extract_text_from_txt(file_path)

        parsed_json = json.loads(raw_text)

        formatted_json = json.dumps(
            parsed_json,
            indent=2,
            ensure_ascii=False,
        )

        return clean_extracted_text(formatted_json)

    except json.JSONDecodeError as exc:
        raise DocumentExtractionError(f"Invalid JSON document: {exc}") from exc

    except Exception as exc:
        logger.exception(
            "JSON extraction failed for %s",
            file_path,
        )

        raise DocumentExtractionError(f"Unable to extract JSON text: {exc}") from exc


def validate_document_path(
    file_path: str | Path,
) -> Path:
    """
    Validate that the document exists and is supported.
    """

    path = Path(file_path)

    if not path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Document file not found: {path}",
        )

    if not path.is_file():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="The supplied document path is not a file.",
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        allowed_extensions = ", ".join(sorted(SUPPORTED_EXTENSIONS))

        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Unsupported document type '{extension}'. Supported types: {allowed_extensions}."
            ),
        )

    return path


def extract_document_text(
    file_path: str | Path,
) -> str:
    """
    Extract text from a supported document.

    Supported:
    - PDF
    - TXT
    - Markdown
    - DOCX
    - CSV
    - JSON
    """

    path = validate_document_path(file_path)

    extension = path.suffix.lower()

    extractors = {
        ".pdf": extract_text_from_pdf,
        ".txt": extract_text_from_txt,
        ".md": extract_text_from_markdown,
        ".docx": extract_text_from_docx,
        ".csv": extract_text_from_csv,
        ".json": extract_text_from_json,
    }

    extractor = extractors.get(extension)

    if extractor is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(f"No text extractor is available for '{extension}'."),
        )

    try:
        extracted_text = extractor(path)

        if not extracted_text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=(
                    "No readable text was found in the document. "
                    "The file may be empty or contain scanned images."
                ),
            )

        logger.info(
            "Extracted %s characters from %s",
            len(extracted_text),
            path.name,
        )

        return extracted_text

    except HTTPException:
        raise

    except DocumentExtractionError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=str(exc),
        ) from exc

    except Exception as exc:
        logger.exception(
            "Unexpected extraction error for %s",
            path,
        )

        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=("Unexpected error while extracting document text."),
        ) from exc
