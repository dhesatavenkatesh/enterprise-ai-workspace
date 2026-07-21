from __future__ import annotations

import logging
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader


logger = logging.getLogger(__name__)


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".markdown",
}

SUPPORTED_CHUNK_SIZES = {
    256,
    512,
    1024,
}

SUPPORTED_CHUNK_OVERLAPS = {
    20,
    50,
    100,
}


class DocumentProcessingError(Exception):
    """Raised when a document cannot be processed."""


class UnsupportedDocumentTypeError(DocumentProcessingError):
    """Raised when the uploaded file type is unsupported."""


class EmptyDocumentError(DocumentProcessingError):
    """Raised when no readable text can be extracted."""


@dataclass
class ExtractedPage:
    page_number: int
    content: str
    character_count: int


@dataclass
class DocumentChunk:
    chunk_index: int
    content: str
    character_count: int
    word_count: int
    start_position: int
    end_position: int
    metadata: dict[str, Any]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass
class ProcessedDocument:
    document_id: str
    file_name: str
    file_path: str
    document_type: str
    full_text: str
    character_count: int
    word_count: int
    page_count: int
    chunk_size: int
    chunk_overlap: int
    chunks: list[DocumentChunk]
    metadata: dict[str, Any]

    @property
    def chunk_count(self) -> int:
        return len(self.chunks)

    def to_dict(self) -> dict[str, Any]:
        data = asdict(self)
        data["chunk_count"] = self.chunk_count
        return data


class DocumentProcessor:
    """
    Extracts, cleans and chunks enterprise knowledge documents.

    Supported file types:
    - PDF
    - DOCX
    - TXT
    - Markdown
    """

    def __init__(
        self,
        chunk_size: int = 512,
        chunk_overlap: int = 50,
    ) -> None:
        self.chunk_size = self._validate_chunk_size(chunk_size)
        self.chunk_overlap = self._validate_chunk_overlap(chunk_overlap)

        if self.chunk_overlap >= self.chunk_size:
            raise ValueError(
                "chunk_overlap must be smaller than chunk_size."
            )

    def process_document(
        self,
        *,
        document_id: str,
        file_path: str | Path,
        file_name: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> ProcessedDocument:
        """
        Complete processing workflow:

        Read file
        ↓
        Extract text
        ↓
        Clean text
        ↓
        Chunk text
        ↓
        Attach metadata
        """

        path = Path(file_path).resolve()

        if not path.exists():
            raise DocumentProcessingError(
                f"Document file was not found: {path}"
            )

        if not path.is_file():
            raise DocumentProcessingError(
                f"Document path is not a file: {path}"
            )

        extension = path.suffix.lower()

        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentTypeError(
                f"Unsupported document type: {extension}"
            )

        resolved_file_name = file_name or path.name
        base_metadata = metadata.copy() if metadata else {}

        base_metadata.update(
            {
                "document_id": str(document_id),
                "file_name": resolved_file_name,
                "file_path": str(path),
                "document_type": self._get_document_type(extension),
            }
        )

        extracted_pages = self.extract_text(path)

        cleaned_pages: list[ExtractedPage] = []

        for page in extracted_pages:
            cleaned_content = self.clean_text(page.content)

            if not cleaned_content:
                continue

            cleaned_pages.append(
                ExtractedPage(
                    page_number=page.page_number,
                    content=cleaned_content,
                    character_count=len(cleaned_content),
                )
            )

        if not cleaned_pages:
            raise EmptyDocumentError(
                f"No readable text was found in {resolved_file_name}."
            )

        full_text = self._combine_pages(cleaned_pages)

        chunks = self.create_chunks(
            text=full_text,
            document_id=str(document_id),
            file_name=resolved_file_name,
            document_type=self._get_document_type(extension),
            metadata=base_metadata,
        )

        if not chunks:
            raise EmptyDocumentError(
                f"No chunks were generated for {resolved_file_name}."
            )

        processed_document = ProcessedDocument(
            document_id=str(document_id),
            file_name=resolved_file_name,
            file_path=str(path),
            document_type=self._get_document_type(extension),
            full_text=full_text,
            character_count=len(full_text),
            word_count=len(full_text.split()),
            page_count=len(cleaned_pages),
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            chunks=chunks,
            metadata=base_metadata,
        )

        logger.info(
            "Document processed successfully: document_id=%s chunks=%s",
            document_id,
            processed_document.chunk_count,
        )

        return processed_document

    def extract_text(
        self,
        file_path: str | Path,
    ) -> list[ExtractedPage]:
        path = Path(file_path)
        extension = path.suffix.lower()

        extractors = {
            ".pdf": self._extract_pdf,
            ".docx": self._extract_docx,
            ".txt": self._extract_text_file,
            ".md": self._extract_text_file,
            ".markdown": self._extract_text_file,
        }

        extractor = extractors.get(extension)

        if extractor is None:
            raise UnsupportedDocumentTypeError(
                f"Unsupported file extension: {extension}"
            )

        try:
            return extractor(path)
        except DocumentProcessingError:
            raise
        except Exception as exc:
            logger.exception(
                "Failed to extract document text: %s",
                path,
            )
            raise DocumentProcessingError(
                f"Unable to extract text from {path.name}: {exc}"
            ) from exc

    def clean_text(self, text: str) -> str:
        """
        Cleans extracted text while preserving paragraph structure.
        """

        if not text:
            return ""

        cleaned = text.replace("\x00", " ")
        cleaned = cleaned.replace("\r\n", "\n")
        cleaned = cleaned.replace("\r", "\n")

        cleaned = re.sub(
            r"[\u200b\u200c\u200d\ufeff]",
            "",
            cleaned,
        )

        cleaned = re.sub(
            r"[ \t]+",
            " ",
            cleaned,
        )

        cleaned = re.sub(
            r" *\n *",
            "\n",
            cleaned,
        )

        cleaned = re.sub(
            r"\n{3,}",
            "\n\n",
            cleaned,
        )

        cleaned = re.sub(
            r"(?<!\n)\n(?!\n)",
            " ",
            cleaned,
        )

        cleaned = re.sub(
            r" {2,}",
            " ",
            cleaned,
        )

        cleaned = re.sub(
            r"\n{3,}",
            "\n\n",
            cleaned,
        )

        return cleaned.strip()

    def create_chunks(
        self,
        *,
        text: str,
        document_id: str,
        file_name: str,
        document_type: str,
        metadata: dict[str, Any] | None = None,
    ) -> list[DocumentChunk]:
        """
        Creates overlapping chunks measured approximately by words.

        Example:
        chunk_size=512
        chunk_overlap=50
        """

        cleaned_text = self.clean_text(text)

        if not cleaned_text:
            return []

        words = cleaned_text.split()

        if not words:
            return []

        chunks: list[DocumentChunk] = []
        start_word = 0
        chunk_index = 0
        total_words = len(words)
        search_position = 0

        while start_word < total_words:
            end_word = min(
                start_word + self.chunk_size,
                total_words,
            )

            chunk_words = words[start_word:end_word]
            chunk_content = " ".join(chunk_words).strip()

            if not chunk_content:
                break

            start_position = cleaned_text.find(
                chunk_content,
                search_position,
            )

            if start_position == -1:
                start_position = search_position

            end_position = start_position + len(chunk_content)

            chunk_metadata = {
                **(metadata or {}),
                "document_id": str(document_id),
                "file_name": file_name,
                "document_type": document_type,
                "chunk_index": chunk_index,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "word_start": start_word,
                "word_end": end_word,
                "character_start": start_position,
                "character_end": end_position,
            }

            chunks.append(
                DocumentChunk(
                    chunk_index=chunk_index,
                    content=chunk_content,
                    character_count=len(chunk_content),
                    word_count=len(chunk_words),
                    start_position=start_position,
                    end_position=end_position,
                    metadata=chunk_metadata,
                )
            )

            if end_word >= total_words:
                break

            next_start = end_word - self.chunk_overlap

            if next_start <= start_word:
                next_start = end_word

            start_word = next_start
            search_position = max(
                0,
                end_position - 1000,
            )
            chunk_index += 1

        return chunks

    def _extract_pdf(
        self,
        file_path: Path,
    ) -> list[ExtractedPage]:
        reader = PdfReader(str(file_path))
        pages: list[ExtractedPage] = []

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentProcessingError(
                    "Encrypted PDF files are not supported."
                ) from exc

        for page_index, page in enumerate(
            reader.pages,
            start=1,
        ):
            try:
                text = page.extract_text() or ""
            except Exception:
                logger.warning(
                    "Unable to extract PDF page %s from %s",
                    page_index,
                    file_path.name,
                )
                text = ""

            pages.append(
                ExtractedPage(
                    page_number=page_index,
                    content=text,
                    character_count=len(text),
                )
            )

        return pages

    def _extract_docx(
        self,
        file_path: Path,
    ) -> list[ExtractedPage]:
        document = DocxDocument(str(file_path))
        content_blocks: list[str] = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()

            if paragraph_text:
                content_blocks.append(paragraph_text)

        for table in document.tables:
            for row in table.rows:
                row_values = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if row_values:
                    content_blocks.append(
                        " | ".join(row_values)
                    )

        combined_text = "\n\n".join(content_blocks)

        return [
            ExtractedPage(
                page_number=1,
                content=combined_text,
                character_count=len(combined_text),
            )
        ]

    def _extract_text_file(
        self,
        file_path: Path,
    ) -> list[ExtractedPage]:
        encodings = (
            "utf-8",
            "utf-8-sig",
            "utf-16",
            "cp1252",
            "latin-1",
        )

        content: str | None = None
        last_error: Exception | None = None

        for encoding in encodings:
            try:
                content = file_path.read_text(
                    encoding=encoding,
                )
                break
            except UnicodeDecodeError as exc:
                last_error = exc

        if content is None:
            raise DocumentProcessingError(
                f"Unable to decode {file_path.name}: {last_error}"
            )

        return [
            ExtractedPage(
                page_number=1,
                content=content,
                character_count=len(content),
            )
        ]

    @staticmethod
    def _combine_pages(
        pages: list[ExtractedPage],
    ) -> str:
        content_parts: list[str] = []

        for page in pages:
            if page.content.strip():
                content_parts.append(page.content.strip())

        return "\n\n".join(content_parts).strip()

    @staticmethod
    def _get_document_type(extension: str) -> str:
        document_types = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".txt": "txt",
            ".md": "markdown",
            ".markdown": "markdown",
        }

        return document_types.get(
            extension.lower(),
            "unknown",
        )

    @staticmethod
    def _validate_chunk_size(chunk_size: int) -> int:
        if chunk_size not in SUPPORTED_CHUNK_SIZES:
            raise ValueError(
                "chunk_size must be one of: "
                f"{sorted(SUPPORTED_CHUNK_SIZES)}"
            )

        return chunk_size

    @staticmethod
    def _validate_chunk_overlap(
        chunk_overlap: int,
    ) -> int:
        if chunk_overlap not in SUPPORTED_CHUNK_OVERLAPS:
            raise ValueError(
                "chunk_overlap must be one of: "
                f"{sorted(SUPPORTED_CHUNK_OVERLAPS)}"
            )

        return chunk_overlap


def process_document(
    *,
    document_id: str,
    file_path: str | Path,
    file_name: str | None = None,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
    metadata: dict[str, Any] | None = None,
) -> ProcessedDocument:
    """
    Convenience function for processing a document without manually
    creating a DocumentProcessor instance.
    """

    processor = DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )

    return processor.process_document(
        document_id=document_id,
        file_path=file_path,
        file_name=file_name,
        metadata=metadata,
    )